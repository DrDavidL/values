import streamlit as st
from google import genai
import io
from docx import Document

st.set_page_config(page_title="Value Cards App", page_icon="🔮", layout="wide")
st.title("Value Cards App")
st.write("Select your core values from each category below. Toggle the items you resonate with, and your Tier 2 values will update on the right side.")


# Store your password securely in an environment variable or secrets.toml
password_key = st.secrets["app"]["password"]
api_key = st.secrets["GEMINI_API_KEY"]

def check_password(widget_key_suffix="") -> bool:
    """
    Prompt for password entry with improved usability and robust feedback.
    
    Args:
        widget_key_suffix: Suffix to ensure unique Streamlit widget keys.
    Returns:
        bool: True if password is correct, else False.
    """
    # Session state initialization
    if "password_correct" not in st.session_state:
        st.session_state.password_correct = False
    if "login_attempts" not in st.session_state:
        st.session_state.login_attempts = 0

    password_key_name = f"password_{widget_key_suffix}"

    def password_entered():
        entered_password = st.session_state[password_key_name]
        if entered_password == password_key:
            st.session_state.password_correct = True
            st.session_state.login_attempts = 0
            # Optionally: st.success("Login successful!")
        else:
            st.session_state.password_correct = False
            st.session_state.login_attempts += 1
            st.session_state[password_key_name] = ""  # Reset entry field

    if not st.session_state.password_correct:
        st.text_input(
            "Enter access password",
            type="password",
            key=password_key_name,
            on_change=password_entered,
            placeholder="Password",
            help="Contact David Liebovitz, MD for access."
        )

        if st.session_state.login_attempts > 0:
            st.error(
                f"Incorrect password. Attempts: {st.session_state.login_attempts}"
            )

        st.caption(
            "*If you need an updated access password, contact David Liebovitz, MD.*"
        )
        return False
    return True

with st.expander("About this App"):
    st.write(
        "This app helps you identify and organize your core values. \n\n"
        "Select your Tier 2 values from the list, then refine to only **two** Tier 1 values. (Not easy!) \n"
        "Where applicable, add Tier 2 values under your Tier 1 values to generate draft value statements. \n\n"
        "This app was authored by David Liebovitz, MD, and is open-source on [GitHub](https://github.com/DrDavidL/values)."
    )
    st.info("If you have an optional API key for Gemini, AI versions of your value statements are re-phrased using the [Gemini Language Model](https://aistudio.google.com/): ")

    st.markdown(
        "**Sources:** This app’s values list is based on curated items from "
        "[James Clear](https://jamesclear.com/core-values), "
        "[Brené Brown](https://brenebrown.com/resources/dare-to-lead-list-of-values/), "
        "[Nir Eyal](https://www.nirandfar.com/common-values/), and "
        "[Colin Breck](https://blog.colinbreck.com/understanding-our-core-values-an-exercise-for-individuals-and-teams/)."
    )

# Define categories and associated values
categories = {
    "I. Authenticity & Integrity": [
        "Authenticity", "Honesty", "Integrity", "Sincerity", "Vulnerability", "Truth", 
        "Uniqueness", "Intuition", "Self-respect"
    ],
    "II. Personal Growth & Achievement": [
        "Achievement", "Ambition", "Growth", "Learning", "Determination", "Perseverance",
        "Discipline / Self-discipline", "Initiative", "Curiosity", "Innovation", 
        "Creativity", "Resourcefulness", "Accountability", "Legacy", "Risk-taking", 
        "Personal Fulfillment", "Understanding", "Wisdom", "Adaptability", "Commitment", 
        "Competence", "Excellence", "Knowledge", "Pride", "Success", "Diligence"
    ],
    "III. Independence & Freedom": [
        "Autonomy", "Freedom", "Independence", "Self-reliance"
    ],
    "IV. Interpersonal Relationships & Connection": [
        "Connection", "Compassion", "Empathy", "Kindness", "Caring", "Friendship", 
        "Belonging", "Community", "Loyalty", "Family", "Teamwork", "Collaboration", 
        "Cooperation", "Love", "Trust"
    ],
    "V. Leadership & Influence": [
        "Authority", "Leadership", "Influence", "Confidence", "Power", "Recognition", "Vision"
    ],
    "VI. Well-being, Balance & Happiness": [
        "Balance", "Harmony", "Health", "Well-being", "Joy", "Contentment", "Peace", 
        "Serenity", "Fun", "Optimism", "Patience", "Gratitude", "Hope", "Humor", 
        "Leisure", "Wholeheartedness", "Simplicity"
    ],
    "VII. Ethical & Moral Values": [
        "Fairness", "Justice", "Ethics", "Responsibility", "Respect", "Honor", "Humility", 
        "Forgiveness", "Non-violence", "Tradition", "Dignity"
    ],
    "VIII. Stability & Security": [
        "Financial Stability", "Job Security", "Safety", "Security", "Order", "Stability", 
        "Dependability", "Reliability", "Thrift", "Home", "Wealth"
    ],
    "IX. Generosity & Contribution": [
        "Generosity", "Contribution", "Service", "Giving", "Making a Difference", 
        "Stewardship", "Altruism", "Future Generations"
    ],
    "X. Aesthetics & Appreciation": [
        "Beauty", "Appreciation", "Grace"
    ],
    "XI. Exploration & Experience": [
        "Adventure", "Boldness", "Bravery", "Travel", "Challenge", "Courage"
    ],
    "XII. Miscellaneous & Broader Values": [
        "Citizenship", "Efficiency", "Enthusiasm", "Equality", "Faith", "Inclusion", 
        "Career", "Nature", "Environment", "Openness", "Parenting", "Patriotism", 
        "Self-expression", "Sportsmanship", "Time", "Usefulness", "Diversity", "Spirituality"
    ]
}

# Compute the final selection from toggles
final_selection = []

col1, col3 = st.columns(2, gap="medium")
with col1:
    st.write("## Step 1. Select Your Values")
    # Loop over each category and each value within the category.
    for category, values in categories.items():
        with st.expander(category):
            for value in values:
                # Use a toggle for each value.
                if st.toggle(label=value, key=f"{category}_{value}"):
                    final_selection.append(value)

    # Store the final selection in session state if needed later.
    st.session_state.final_selection = final_selection

with col3:
    st.write("## Your Tier 2 Values List")
    if final_selection:
        unique_values = sorted(set(final_selection))
        for value in unique_values:
            st.write(f"- {value}")
    else:
        st.info("No values selected yet.")

    if final_selection:
        st.write("## Step 2: Choose Two Values as your Tier 1 Values")
        unique_values = sorted(set(final_selection))
        selected_cores = st.multiselect(
            "Select two core values from your Tier 2 list as your two Tier 1 Values:",
            options=unique_values, key="core_values", max_selections=2
        )
        
        if len(selected_cores) == 2:
            st.write("## Step 3: Select Tier 2 Values to complement your primary Tier 1 Values")
            st.write("Where it makes sense to you, add Tier 2 values to complement your Tier 1 Values.")
            
            # Filter out Tier 1 values from the Tier 2 selection options
            tier2_options = [value for value in unique_values if value not in selected_cores]
            
            bucket1 = st.multiselect(
                f"Select values for bucket **{selected_cores[0]}**:", 
                options=tier2_options, key="bucket1 choices"
            )
            bucket2 = st.multiselect(
                f"Select values for bucket **{selected_cores[1]}**:", 
                options=tier2_options, key="bucket2 choices"
            )
            
            statement1 = f"I value {selected_cores[0]} supported by " + ", ".join(bucket1)
            statement2 = f"I value {selected_cores[1]} supported by " + ", ".join(bucket2)
            # Display draft statements
            st.write("## Step 4. View Draft Values Statements")
            st.write(f"### 1. {statement1}")
            st.write(f"### 2. {statement2}")
            st.session_state.draft_statements = f"1. {statement1}\n2. {statement2}"
            
            # Save the draft statements and other info into session state for DOCX generation
            
            st.session_state.selected_cores = selected_cores
            st.session_state.bucket1 = bucket1
            st.session_state.bucket2 = bucket2
            ai_version = st.checkbox(
                "Use AI to enhance the draft statements",
                value=False, help="If you have a password, you can use AI to enhance your draft statements.")
            if ai_version:
                if check_password("ai"):
                    if st.button("Draft AI Enhanced Values Statements"):
                        try:
                            client = genai.Client(api_key=api_key)
                                                
                            response = client.models.generate_content(
                                model="gemini-2.0-flash",
                                contents=f"""Rephrase each of the following statements to convey a draft essence of the user\'s core values. Return only numbered rephrased versions, not other comments or information.:
                                Statement 1: {statement1},
                                Statement 2: {statement2}"""
                            )
                            draft_statements = response.text.strip()
                            st.session_state.ai_draft_statements = draft_statements                 

                        except Exception as e:
                            st.error(f"An error occurred during the LLM call: {e}")
                    # else:
                    #     st.info("Please select exactly two core values to create your buckets.")
                    if st.session_state.get("ai_draft_statements"):
                        st.write("### AI Enhanced Statements")
                        st.write(st.session_state.ai_draft_statements)     

    # Provide a download button if draft statements exist
    if "draft_statements" in st.session_state or "ai_draft_statements" in st.session_state:
        def generate_docx():
            document = Document()
            document.add_heading("Value Cards Document", 0)

            # Tier 2 Values Section
            document.add_heading("Tier 2 Values List", level=1)
            if st.session_state.final_selection:
                for val in sorted(set(st.session_state.final_selection)):
                    document.add_paragraph(val, style='List Bullet')
            else:
                document.add_paragraph("No Tier 2 values selected.")

            # Tier 1 Buckets Section
            document.add_heading("Tier 1 Values", level=1)
            selected_cores = st.session_state.get("selected_cores", [])
            bucket1 = st.session_state.get("bucket1", [])
            bucket2 = st.session_state.get("bucket2", [])
            if len(selected_cores) == 2:
                document.add_heading(f"{selected_cores[0]} with:", level=2)
                if bucket1:
                    for val in bucket1:
                        document.add_paragraph(val, style='List Bullet')
                else:
                    document.add_paragraph("No values assigned.")
                
                document.add_heading(f"{selected_cores[1]} with:", level=2)
                if bucket2:
                    for val in bucket2:
                        document.add_paragraph(val, style='List Bullet')
                else:
                    document.add_paragraph("No values assigned.")
            else:
                document.add_paragraph("Tier 1 bucket details not available.")

            # Draft Statements Section
            document.add_heading("Draft Statements", level=1)
            draft = st.session_state.get("draft_statements", "No draft statements generated.")
            document.add_paragraph(draft)

            ai = st.session_state.get("ai_draft_statements", "No AI draft statements generated.")
            document.add_heading("AI Enhanced Draft Statements", level=2)
            document.add_paragraph(ai)                      
            
            # Save the document to a BytesIO buffer and return
            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return buffer

        st.write("## Step 5. Download to Save and Finalize Your Draft Statements")
        docx_buffer = generate_docx()
        st.download_button(
            label="Download your DOCX file for final editing!",
            data=docx_buffer,
            file_name="Value_Cards_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
